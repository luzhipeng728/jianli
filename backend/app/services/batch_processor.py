"""批量简历处理器 - 后台异步处理"""
import asyncio
from typing import AsyncGenerator
from app.services.task_manager import TaskManager, TaskStatus, task_manager, BatchTask
from app.services.resume_parser import ResumeParser
from app.services.file_processor import FileType, FileProcessor


class BatchProcessor:
    """批量简历处理器"""

    def __init__(self):
        self.task_manager = task_manager
        self.resume_parser = ResumeParser()
        self.file_processor = FileProcessor()
        self._running_tasks: dict[str, asyncio.Task] = {}

    async def process_batch(
        self,
        batch_id: str,
        files: list[tuple[str, bytes]]  # [(filename, content), ...]
    ) -> AsyncGenerator[dict, None]:
        """处理批量任务，流式返回进度

        Args:
            batch_id: 批量任务ID
            files: 文件列表 [(文件名, 文件内容), ...]

        Yields:
            进度更新消息
        """
        batch = await self.task_manager.get_batch(batch_id)
        if not batch:
            yield {"type": "error", "message": "批量任务不存在"}
            return

        # 更新状态为处理中
        batch.status = TaskStatus.PROCESSING
        await self.task_manager.update_batch(batch)

        yield {
            "type": "started",
            "batch_id": batch_id,
            "total": batch.total_files,
            "message": f"开始处理 {batch.total_files} 份简历"
        }

        # 创建文件映射
        file_map = {name: content for name, content in files}

        # 并发处理（控制并发数）
        semaphore = asyncio.Semaphore(5)  # 最多5个并发

        async def process_single(file_task):
            async with semaphore:
                return await self._process_file(batch_id, file_task, file_map)

        # 按批次处理，每处理完一个就汇报进度
        for file_task in batch.files:
            # 跳过已成功的
            if file_task.status == TaskStatus.SUCCESS:
                continue

            # 检查是否是重试任务
            if file_task.status not in [TaskStatus.PENDING, TaskStatus.RETRYING]:
                continue

            filename = file_task.file_name
            content = file_map.get(filename)

            if not content:
                await self.task_manager.update_file_status(
                    batch_id, file_task.file_id,
                    TaskStatus.FAILED,
                    error="文件内容缺失"
                )
                yield {
                    "type": "file_failed",
                    "file_id": file_task.file_id,
                    "file_name": filename,
                    "error": "文件内容缺失"
                }
                continue

            # 更新为处理中
            await self.task_manager.update_file_status(
                batch_id, file_task.file_id,
                TaskStatus.PROCESSING,
                progress=10
            )

            yield {
                "type": "file_processing",
                "file_id": file_task.file_id,
                "file_name": filename,
                "progress": 10
            }

            try:
                # 使用流式解析简历
                file_type, text = self.file_processor.process_file(filename, content)

                # 如果是图片或扫描件PDF，传递原始数据用于OCR
                image_data = content if file_type in (FileType.IMAGE, FileType.PDF_SCANNED) else None

                resume = None
                # 传递raw_content用于存储原始文件
                async for item in self.resume_parser.parse_with_text_stream(filename, file_type, text, image_data, raw_content=content):
                    if item["type"] == "chunk":
                        # 流式返回LLM输出片段
                        yield {
                            "type": "llm_chunk",
                            "file_id": file_task.file_id,
                            "file_name": filename,
                            "content": item["content"]
                        }
                    elif item["type"] == "status":
                        # OCR状态更新
                        yield {
                            "type": "file_status",
                            "file_id": file_task.file_id,
                            "file_name": filename,
                            "message": item["message"]
                        }
                    elif item["type"] == "done":
                        resume = item["data"]

                if not resume:
                    raise Exception("解析失败：未能获取简历数据")

                # 进行维度分析
                yield {
                    "type": "file_status",
                    "file_id": file_task.file_id,
                    "file_name": filename,
                    "message": "正在进行AI多维度分析..."
                }

                try:
                    from app.services.dimension_analyzer import analyze_resume_dimensions
                    analysis_result = await analyze_resume_dimensions(
                        resume.model_dump(mode="json"),
                        analysis_type="screening"
                    )

                    # 将维度分析结果添加到简历数据
                    from app.models.resume import DimensionAnalysis, DimensionScore
                    resume.dimension_analysis = DimensionAnalysis(
                        dimensions=[
                            DimensionScore(
                                name=d.name,
                                score=d.score,
                                level=d.level,
                                highlights=d.highlights,
                                concerns=d.concerns
                            )
                            for d in analysis_result.dimensions
                        ],
                        overall_score=analysis_result.overall_score,
                        summary=analysis_result.summary,
                        recommendations=analysis_result.recommendations,
                        analysis_date=await self._get_current_time()
                    )

                    # 更新ES中的简历数据
                    from app.services.es_client import ESClient
                    es_client = ESClient()
                    doc = resume.model_dump(mode="json")
                    es_client.index_document("resumes", resume.id, doc)

                except Exception as analysis_error:
                    print(f"[BatchProcessor] 维度分析失败: {analysis_error}")
                    # 维度分析失败不影响整体流程

                # 构建详细结果
                result_data = {
                    "resume_id": resume.id,
                    "name": resume.basic_info.name,
                    "phone": resume.basic_info.phone,
                    "email": resume.basic_info.email,
                    "skills": resume.skills.hard_skills[:10] if resume.skills.hard_skills else [],
                    "warnings_count": len(resume.warnings),
                    "warnings": [{"type": w.type, "message": w.message} for w in resume.warnings[:5]]
                }

                # 添加维度分析结果到返回数据
                if resume.dimension_analysis:
                    result_data["dimension_analysis"] = {
                        "overall_score": resume.dimension_analysis.overall_score,
                        "summary": resume.dimension_analysis.summary,
                        "dimensions": [
                            {"name": d.name, "score": d.score, "level": d.level}
                            for d in resume.dimension_analysis.dimensions
                        ]
                    }

                # 更新为成功
                await self.task_manager.update_file_status(
                    batch_id, file_task.file_id,
                    TaskStatus.SUCCESS,
                    progress=100,
                    result=result_data
                )

                yield {
                    "type": "file_success",
                    "file_id": file_task.file_id,
                    "file_name": filename,
                    "resume_id": resume.id,
                    "name": resume.basic_info.name,
                    "phone": resume.basic_info.phone,
                    "email": resume.basic_info.email,
                    "skills": result_data["skills"],
                    "warnings_count": result_data["warnings_count"],
                    "warnings": result_data["warnings"]
                }

            except Exception as e:
                error_msg = str(e)

                # 检查是否可以重试
                batch = await self.task_manager.get_batch(batch_id)
                current_task = next((f for f in batch.files if f.file_id == file_task.file_id), None)

                if current_task and current_task.retry_count < self.task_manager.MAX_RETRIES:
                    # 自动重试 - 使用OCR方式
                    await asyncio.sleep(self.task_manager.RETRY_DELAY)
                    await self.task_manager.retry_file(batch_id, file_task.file_id)

                    yield {
                        "type": "file_retrying",
                        "file_id": file_task.file_id,
                        "file_name": filename,
                        "retry_count": current_task.retry_count + 1,
                        "error": error_msg,
                        "message": "尝试使用OCR方式重新解析..."
                    }

                    # 使用OCR方式重试
                    try:
                        yield {
                            "type": "file_status",
                            "file_id": file_task.file_id,
                            "file_name": filename,
                            "message": "正在进行OCR识别..."
                        }

                        # 尝试将文件转为图片进行OCR
                        ocr_text = await self._try_ocr_fallback(filename, content)

                        if ocr_text:
                            yield {
                                "type": "file_status",
                                "file_id": file_task.file_id,
                                "file_name": filename,
                                "message": f"OCR完成，提取到{len(ocr_text)}字符，正在解析..."
                            }

                            # 使用OCR文本重新解析
                            resume = None
                            async for item in self.resume_parser.parse_with_text_stream(
                                filename, FileType.IMAGE, ocr_text, None
                            ):
                                if item["type"] == "chunk":
                                    yield {
                                        "type": "llm_chunk",
                                        "file_id": file_task.file_id,
                                        "file_name": filename,
                                        "content": item["content"]
                                    }
                                elif item["type"] == "done":
                                    resume = item["data"]

                            if resume:
                                result_data = {
                                    "resume_id": resume.id,
                                    "name": resume.basic_info.name,
                                    "phone": resume.basic_info.phone,
                                    "email": resume.basic_info.email,
                                    "skills": resume.skills.hard_skills[:10] if resume.skills.hard_skills else [],
                                    "warnings_count": len(resume.warnings),
                                    "warnings": [{"type": w.type, "message": w.message} for w in resume.warnings[:5]]
                                }
                                await self.task_manager.update_file_status(
                                    batch_id, file_task.file_id,
                                    TaskStatus.SUCCESS,
                                    progress=100,
                                    result=result_data
                                )
                                yield {
                                    "type": "file_success",
                                    "file_id": file_task.file_id,
                                    "file_name": filename,
                                    "resume_id": resume.id,
                                    "name": resume.basic_info.name,
                                    "phone": resume.basic_info.phone,
                                    "email": resume.basic_info.email,
                                    "skills": result_data["skills"],
                                    "warnings_count": result_data["warnings_count"],
                                    "warnings": result_data["warnings"]
                                }
                            else:
                                raise Exception("OCR解析后未能提取简历数据")
                        else:
                            raise Exception("OCR识别失败")

                    except Exception as retry_error:
                        await self.task_manager.update_file_status(
                            batch_id, file_task.file_id,
                            TaskStatus.FAILED,
                            error=str(retry_error)
                        )
                        yield {
                            "type": "file_failed",
                            "file_id": file_task.file_id,
                            "file_name": filename,
                            "error": str(retry_error)
                        }
                else:
                    await self.task_manager.update_file_status(
                        batch_id, file_task.file_id,
                        TaskStatus.FAILED,
                        error=error_msg
                    )

                    yield {
                        "type": "file_failed",
                        "file_id": file_task.file_id,
                        "file_name": filename,
                        "error": error_msg
                    }

            # 获取最新进度
            batch = await self.task_manager.get_batch(batch_id)
            yield {
                "type": "progress",
                "batch_id": batch_id,
                "completed": batch.completed,
                "failed": batch.failed,
                "total": batch.total_files,
                "progress": batch.progress
            }

        # 完成
        batch = await self.task_manager.get_batch(batch_id)
        yield {
            "type": "completed",
            "batch_id": batch_id,
            "status": batch.status.value,
            "total": batch.total_files,
            "success": batch.completed,
            "failed": batch.failed,
            "message": f"处理完成: {batch.completed} 成功, {batch.failed} 失败"
        }

    async def _process_file(
        self,
        batch_id: str,
        file_task,
        file_map: dict[str, bytes]
    ) -> dict:
        """处理单个文件"""
        filename = file_task.file_name
        content = file_map.get(filename)

        if not content:
            return {"status": "failed", "error": "文件内容缺失"}

        try:
            resume = await self.resume_parser.parse(filename, content)
            return {
                "status": "success",
                "resume_id": resume.id,
                "name": resume.basic_info.name
            }
        except Exception as e:
            return {"status": "failed", "error": str(e)}

    async def _get_current_time(self) -> str:
        """获取当前时间字符串"""
        from datetime import datetime
        return datetime.now().isoformat()

    async def _try_ocr_fallback(self, filename: str, content: bytes) -> str | None:
        """尝试使用OCR方式提取文本

        对于解析失败的文件，尝试转换为图片后进行OCR识别
        """
        from app.services.llm_client import LLMClient

        llm_client = LLMClient()
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        try:
            if ext == 'pdf':
                # PDF转图片后OCR
                images = self.file_processor.extract_images_from_pdf(content)
                if not images:
                    return None

                ocr_texts = []
                for i, img_data in enumerate(images[:5]):  # 最多处理5页
                    ocr_text = await llm_client.ocr(img_data, f"page_{i+1}.png")
                    if ocr_text:
                        ocr_texts.append(ocr_text)

                return "\n\n".join(ocr_texts) if ocr_texts else None

            elif ext in ('jpg', 'jpeg', 'png', 'bmp', 'webp'):
                # 图片直接OCR
                return await llm_client.ocr(content, filename)

            elif ext == 'docx':
                # Word文档 - 尝试用python-docx重新提取
                try:
                    import io
                    from docx import Document
                    doc = Document(io.BytesIO(content))
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text.strip())
                    # 也提取表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                            if row_text:
                                paragraphs.append(row_text)
                    return "\n".join(paragraphs) if paragraphs else None
                except Exception:
                    return None

            elif ext == 'doc':
                # 老版Word文档，需要转换工具，暂不支持OCR回退
                return None

            else:
                return None

        except Exception as e:
            print(f"OCR fallback failed for {filename}: {e}")
            return None


# 全局处理器实例
batch_processor = BatchProcessor()
