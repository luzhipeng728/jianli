"""后台任务处理 Worker - 支持最高10并发"""
import asyncio
from typing import Optional
from app.services.task_manager import TaskManager, TaskStatus, task_manager
from app.services.resume_parser import ResumeParser
from app.services.file_processor import FileProcessor, FileType


class BackgroundWorker:
    """后台简历处理 Worker"""

    MAX_CONCURRENCY = 10  # 最大并发数（太高会导致API响应变慢）
    QUEUE_CHECK_INTERVAL = 2  # 队列检查间隔（秒）
    FILE_PROCESS_TIMEOUT = 180  # 单个文件处理超时（秒）

    def __init__(self):
        self.task_manager = task_manager
        self.resume_parser = ResumeParser()
        self.file_processor = FileProcessor()
        self._running = False
        self._worker_task: Optional[asyncio.Task] = None
        self._semaphore = asyncio.Semaphore(self.MAX_CONCURRENCY)
        self._active_tasks: dict[str, asyncio.Task] = {}

    async def start(self):
        """启动 Worker"""
        if self._running:
            return
        self._running = True

        # 恢复卡住的任务
        await self._recover_stuck_tasks()

        self._worker_task = asyncio.create_task(self._run_loop())
        print(f"[Worker] 后台处理器已启动，最大并发: {self.MAX_CONCURRENCY}")

    async def _recover_stuck_tasks(self):
        """恢复未完成的任务（processing 或 queued 状态）"""
        batches = await self.task_manager.get_all_batches()
        recovered = 0

        for batch in batches:
            # 检查是否有未完成的文件
            pending_files = [f for f in batch.files if f.status in [TaskStatus.PROCESSING, TaskStatus.QUEUED, TaskStatus.PENDING]]

            if pending_files and batch.status != TaskStatus.SUCCESS:
                # 重置 processing 状态为 queued
                for file_task in batch.files:
                    if file_task.status == TaskStatus.PROCESSING:
                        file_task.status = TaskStatus.QUEUED
                        file_task.progress = 0
                        recovered += 1

                # 更新批量任务状态并重新入队
                batch.status = TaskStatus.PROCESSING
                await self.task_manager.update_batch(batch)
                await self.task_manager.enqueue_batch(batch.batch_id)
                print(f"[Worker] 恢复批次 {batch.batch_id[:8]}...，{len(pending_files)} 个待处理文件")

        if recovered > 0:
            print(f"[Worker] 重置了 {recovered} 个 processing 状态的任务")

    async def stop(self):
        """停止 Worker"""
        self._running = False
        if self._worker_task:
            self._worker_task.cancel()
            try:
                await self._worker_task
            except asyncio.CancelledError:
                pass
        # 取消所有活跃任务
        for task in self._active_tasks.values():
            task.cancel()
        print("[Worker] 后台处理器已停止")

    async def _run_loop(self):
        """主循环 - 监听队列并处理任务"""
        while self._running:
            try:
                # 从队列获取批量任务
                batch_id = await self.task_manager.dequeue_batch(timeout=self.QUEUE_CHECK_INTERVAL)
                if batch_id:
                    # 启动批量处理（不等待完成）
                    asyncio.create_task(self._process_batch(batch_id))
            except asyncio.CancelledError:
                break
            except Exception as e:
                print(f"[Worker] 队列处理错误: {e}")
                await asyncio.sleep(1)

    async def _process_batch(self, batch_id: str):
        """处理批量任务"""
        batch = await self.task_manager.get_batch(batch_id)
        if not batch:
            print(f"[Worker] 批量任务 {batch_id} 不存在")
            return

        print(f"[Worker] 开始处理批量任务 {batch_id}，共 {batch.total_files} 个文件")

        # 更新状态为处理中
        batch.status = TaskStatus.PROCESSING
        await self.task_manager.update_batch(batch)

        # 创建所有文件处理任务
        tasks = []
        for file_task in batch.files:
            # 只处理待处理或重试中的文件
            if file_task.status in [TaskStatus.QUEUED, TaskStatus.PENDING, TaskStatus.RETRYING]:
                task = asyncio.create_task(
                    self._process_file_with_semaphore(batch_id, file_task.file_id, file_task.file_name)
                )
                tasks.append(task)
                self._active_tasks[file_task.file_id] = task

        # 等待所有任务完成
        if tasks:
            await asyncio.gather(*tasks, return_exceptions=True)

        # 清理活跃任务
        for file_task in batch.files:
            self._active_tasks.pop(file_task.file_id, None)

        # 获取最终状态
        batch = await self.task_manager.get_batch(batch_id)
        print(f"[Worker] 批量任务 {batch_id} 完成: {batch.completed} 成功, {batch.failed} 失败")

    async def _process_file_with_semaphore(self, batch_id: str, file_id: str, file_name: str):
        """使用信号量控制并发处理单个文件，带超时保护"""
        async with self._semaphore:
            try:
                await asyncio.wait_for(
                    self._process_file(batch_id, file_id, file_name),
                    timeout=self.FILE_PROCESS_TIMEOUT
                )
            except asyncio.TimeoutError:
                print(f"[Worker] 文件处理超时 ({self.FILE_PROCESS_TIMEOUT}秒): {file_name}")
                await self.task_manager.update_file_status(
                    batch_id, file_id,
                    TaskStatus.FAILED,
                    error=f"处理超时 ({self.FILE_PROCESS_TIMEOUT}秒)"
                )

    async def _process_file(self, batch_id: str, file_id: str, file_name: str):
        """处理单个文件"""
        print(f"[Worker] 开始处理文件: {file_name}")

        # 更新状态为处理中
        await self.task_manager.update_file_status(
            batch_id, file_id,
            TaskStatus.PROCESSING,
            progress=5,
            status_detail="读取文件"
        )

        try:
            # 获取文件内容
            content = await self.task_manager.get_file_content(batch_id, file_id)
            if not content:
                raise Exception("文件内容已过期或不存在")

            # 提取文本
            await self.task_manager.update_file_status(
                batch_id, file_id,
                TaskStatus.PROCESSING,
                progress=10,
                status_detail="提取文本"
            )
            file_type, text = self.file_processor.process_file(file_name, content)

            # 如果是图片或扫描件PDF，传递原始数据用于OCR
            image_data = content if file_type in (FileType.IMAGE, FileType.PDF_SCANNED) else None

            # 流式解析简历
            resume = None
            llm_output = ""

            async for item in self.resume_parser.parse_with_text_stream(
                file_name, file_type, text, image_data, raw_content=content
            ):
                if item["type"] == "chunk":
                    llm_output += item["content"]
                    # 定期更新进度和LLM输出（每50个字符更新一次）
                    if len(llm_output) % 50 == 0:
                        await self.task_manager.update_file_status(
                            batch_id, file_id,
                            TaskStatus.PROCESSING,
                            progress=min(50 + len(llm_output) // 100, 80),
                            llm_output=llm_output,
                            status_detail="LLM解析中"
                        )
                elif item["type"] == "status":
                    # OCR/向量状态更新
                    msg = item['message']
                    if "OCR" in msg:
                        detail = "OCR识别"
                        progress = 20
                    elif "向量" in msg:
                        detail = "创建向量"
                        progress = 85
                    else:
                        detail = msg
                        progress = 30
                    await self.task_manager.update_file_status(
                        batch_id, file_id,
                        TaskStatus.PROCESSING,
                        progress=progress,
                        llm_output=f"[{detail}] {msg}\n{llm_output}",
                        status_detail=detail
                    )
                elif item["type"] == "done":
                    resume = item["data"]

            if not resume:
                raise Exception("解析失败：未能获取简历数据")

            # 保存数据
            await self.task_manager.update_file_status(
                batch_id, file_id,
                TaskStatus.PROCESSING,
                progress=95,
                status_detail="保存数据"
            )

            # 构建结果
            result_data = {
                "resume_id": resume.id,
                "name": resume.basic_info.name,
                "phone": resume.basic_info.phone,
                "email": resume.basic_info.email,
                "skills": resume.skills.hard_skills[:10] if resume.skills.hard_skills else [],
                "warnings_count": len(resume.warnings),
                "warnings": [{"type": w.type, "message": w.message} for w in resume.warnings[:5]]
            }

            # 更新为成功
            await self.task_manager.update_file_status(
                batch_id, file_id,
                TaskStatus.SUCCESS,
                progress=100,
                result=result_data,
                llm_output=llm_output,
                status_detail="完成"
            )

            print(f"[Worker] 文件处理成功: {file_name} -> {resume.basic_info.name}")

        except Exception as e:
            error_msg = str(e)
            print(f"[Worker] 文件处理失败: {file_name} - {error_msg}")

            # 检查是否可以重试
            batch = await self.task_manager.get_batch(batch_id)
            file_task = next((f for f in batch.files if f.file_id == file_id), None)

            if file_task and file_task.retry_count < self.task_manager.MAX_RETRIES:
                # 尝试OCR回退
                retry_success = await self._try_ocr_fallback(batch_id, file_id, file_name)
                if not retry_success:
                    await self.task_manager.update_file_status(
                        batch_id, file_id,
                        TaskStatus.FAILED,
                        error=error_msg
                    )
            else:
                await self.task_manager.update_file_status(
                    batch_id, file_id,
                    TaskStatus.FAILED,
                    error=error_msg
                )

    async def _try_ocr_fallback(self, batch_id: str, file_id: str, file_name: str) -> bool:
        """尝试使用OCR方式重新解析"""
        from app.services.llm_client import LLMClient

        try:
            content = await self.task_manager.get_file_content(batch_id, file_id)
            if not content:
                return False

            llm_client = LLMClient()
            ext = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''

            ocr_text = None

            if ext == 'pdf':
                # PDF转图片后OCR
                images = self.file_processor.extract_images_from_pdf(content)
                if images:
                    ocr_texts = []
                    for i, img_data in enumerate(images[:5]):
                        text = await llm_client.ocr(img_data, f"page_{i+1}.png")
                        if text:
                            ocr_texts.append(text)
                    ocr_text = "\n\n".join(ocr_texts) if ocr_texts else None

            elif ext in ('jpg', 'jpeg', 'png', 'bmp', 'webp'):
                ocr_text = await llm_client.ocr(content, file_name)

            elif ext == 'docx':
                # Word文档重新提取
                try:
                    import io
                    from docx import Document
                    doc = Document(io.BytesIO(content))
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text.strip())
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                            if row_text:
                                paragraphs.append(row_text)
                    ocr_text = "\n".join(paragraphs) if paragraphs else None
                except Exception:
                    pass

            if not ocr_text:
                return False

            # 使用OCR文本重新解析
            resume = None
            llm_output = f"[OCR回退] 提取到 {len(ocr_text)} 字符\n"

            async for item in self.resume_parser.parse_with_text_stream(
                file_name, FileType.IMAGE, ocr_text, None, raw_content=content
            ):
                if item["type"] == "chunk":
                    llm_output += item["content"]
                elif item["type"] == "done":
                    resume = item["data"]

            if resume:
                result_data = {
                    "resume_id": resume.id,
                    "name": resume.basic_info.name,
                    "phone": resume.basic_info.phone,
                    "email": resume.basic_info.email,
                    "skills": resume.skills.hard_skills[:10] if resume.skills.hard_skills else [],
                    "warnings_count": len(resume.warnings),
                    "warnings": [{"type": w.type, "message": w.message} for w in resume.warnings[:5]]
                }

                await self.task_manager.update_file_status(
                    batch_id, file_id,
                    TaskStatus.SUCCESS,
                    progress=100,
                    result=result_data,
                    llm_output=llm_output
                )
                print(f"[Worker] OCR回退成功: {file_name}")
                return True

            return False

        except Exception as e:
            print(f"[Worker] OCR回退失败: {file_name} - {e}")
            return False


# 全局 Worker 实例
background_worker = BackgroundWorker()
