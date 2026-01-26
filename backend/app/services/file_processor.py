import io
import chardet
import pdfplumber
from docx import Document
from PIL import Image
from pathlib import Path
from enum import Enum
from typing import Tuple, List, Optional

class FileType(str, Enum):
    PDF = "pdf"
    PDF_SCANNED = "pdf_scanned"  # 扫描件PDF
    DOCX = "docx"
    TXT = "txt"
    IMAGE = "image"
    UNKNOWN = "unknown"


class FileProcessor:
    SUPPORTED_IMAGES = {".jpg", ".jpeg", ".png", ".bmp", ".webp"}

    # 扫描件检测阈值
    MIN_TEXT_LENGTH = 100  # 少于100字符认为可能是扫描件
    MIN_TEXT_DENSITY = 0.1  # 每页平均字符数过少

    def detect_file_type(self, filename: str) -> FileType:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return FileType.PDF
        elif suffix in {".docx", ".doc"}:
            return FileType.DOCX
        elif suffix == ".txt":
            return FileType.TXT
        elif suffix in self.SUPPORTED_IMAGES:
            return FileType.IMAGE
        return FileType.UNKNOWN

    def extract_text_from_pdf(self, file_content: bytes) -> Tuple[str, bool]:
        """从PDF提取文本，返回 (文本, 是否为扫描件)"""
        text_parts = []
        page_count = 0
        is_scanned = False

        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                page_count = len(pdf.pages)

                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            full_text = "\n".join(text_parts)

            # 检测是否为扫描件
            if page_count > 0:
                avg_chars_per_page = len(full_text) / page_count
                # 如果文本太少，可能是扫描件
                if len(full_text) < self.MIN_TEXT_LENGTH or avg_chars_per_page < 50:
                    is_scanned = True

            return full_text, is_scanned

        except Exception as e:
            print(f"PDF extraction error: {e}")
            return "", True  # 出错时假设为扫描件

    def extract_images_from_pdf(self, file_content: bytes) -> List[bytes]:
        """从PDF中提取图片（用于扫描件OCR）"""
        images = []

        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    # 将页面转为图片
                    pil_image = page.to_image(resolution=200).original
                    img_buffer = io.BytesIO()
                    pil_image.save(img_buffer, format='PNG')
                    images.append(img_buffer.getvalue())
        except Exception as e:
            print(f"PDF image extraction error: {e}")

        return images

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """从DOCX提取文本"""
        try:
            doc = Document(io.BytesIO(file_content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # 也提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())

            return "\n".join(paragraphs)
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ""

    def extract_text_from_txt(self, file_content: bytes) -> str:
        """从TXT提取文本，自动检测编码"""
        try:
            detected = chardet.detect(file_content)
            encoding = detected.get("encoding", "utf-8") or "utf-8"
            confidence = detected.get("confidence", 0)

            # 如果检测置信度低，尝试多种编码
            if confidence < 0.7:
                encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'utf-16']
                for enc in encodings:
                    try:
                        return file_content.decode(enc)
                    except:
                        continue

            return file_content.decode(encoding, errors='replace')
        except Exception as e:
            print(f"TXT extraction error: {e}")
            return file_content.decode('utf-8', errors='replace')

    def process_file(self, filename: str, content: bytes) -> Tuple[FileType, str]:
        """处理文件，返回 (文件类型, 提取的文本)

        对于扫描件PDF，返回 FileType.PDF_SCANNED 和空文本，
        上层需要调用OCR处理。
        """
        file_type = self.detect_file_type(filename)

        # 检查空文件
        if len(content) == 0:
            raise ValueError("文件内容为空")

        if file_type == FileType.PDF:
            text, is_scanned = self.extract_text_from_pdf(content)
            if is_scanned and len(text.strip()) < self.MIN_TEXT_LENGTH:
                # 标记为扫描件，需要OCR
                return FileType.PDF_SCANNED, text
            return file_type, text

        elif file_type == FileType.DOCX:
            text = self.extract_text_from_docx(content)
            return file_type, text

        elif file_type == FileType.TXT:
            text = self.extract_text_from_txt(content)
            return file_type, text

        elif file_type == FileType.IMAGE:
            # 图片需要OCR，返回空文本
            return file_type, ""

        else:
            raise ValueError(f"不支持的文件格式: {filename}")

    def validate_file(self, filename: str, content: bytes) -> Optional[str]:
        """验证文件，返回错误信息或None"""
        if len(content) == 0:
            return "文件内容为空"

        if len(content) > 10 * 1024 * 1024:
            return "文件大小超过10MB限制"

        file_type = self.detect_file_type(filename)
        if file_type == FileType.UNKNOWN:
            return f"不支持的文件格式: {filename}"

        return None
